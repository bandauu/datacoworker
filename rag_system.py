"""
RAG (Retrieval-Augmented Generation) System
Document indexing and semantic search for Data Coworker
"""
import os
import json
from typing import List, Dict
from pathlib import Path

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None


class DocumentRAG:
    """
    Document Retrieval-Augmented Generation system.
    Handles document upload, chunking, embedding, and semantic search.
    """
    
    def __init__(self, persist_dir: str = "database/vector_store"):
        """
        Initialize RAG system with vector store.
        
        Args:
            persist_dir: Directory to persist vector embeddings
        """
        self.persist_dir = persist_dir
        self.embeddings = None
        self.vectorstore = None
        self.documents_metadata = []
        
        # Initialize embeddings model
        self._init_embeddings()
        
        # Load or create vector store
        self._load_or_create_store()
    
    def _init_embeddings(self):
        """Initialize HuggingFace embeddings model"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="all-MiniLM-L6-v2",  # Fast, 80MB model
                cache_folder="./models",
                model_kwargs={'device': 'cpu'}
            )
            print("✅ Embeddings model loaded")
        except Exception as e:
            print(f"❌ Error loading embeddings: {e}")
            self.embeddings = None
    
    def _load_or_create_store(self):
        """Load existing vector store or create new empty one"""
        if not self.embeddings:
            print("⚠️ Embeddings not available, RAG disabled")
            return
        
        os.makedirs(self.persist_dir, exist_ok=True)
        index_path = os.path.join(self.persist_dir, "index.faiss")
        
        if os.path.exists(index_path):
            try:
                self.vectorstore = FAISS.load_local(
                    self.persist_dir, 
                    self.embeddings,
                    allow_dangerous_deserialization=True  # Required for FAISS
                )
                
                # Load metadata
                metadata_path = os.path.join(self.persist_dir, "metadata.json")
                if os.path.exists(metadata_path):
                    with open(metadata_path, 'r') as f:
                        self.documents_metadata = json.load(f)
                
                print(f"✅ Loaded vector store with {len(self.documents_metadata)} documents")
            except Exception as e:
                print(f"⚠️ Error loading vector store: {e}")
                self._create_empty_store()
        else:
            self._create_empty_store()
    
    def _create_empty_store(self):
        """Create empty vector store"""
        try:
            # Create with dummy document
            dummy_doc = Document(
                page_content="Initial document for vector store initialization",
                metadata={"source": "system", "chunk": 0}
            )
            self.vectorstore = FAISS.from_documents([dummy_doc], self.embeddings)
            self.vectorstore.save_local(self.persist_dir)
            print("✅ Created new vector store")
        except Exception as e:
            print(f"❌ Error creating vector store: {e}")
            self.vectorstore = None
    
    def process_pdf(self, file_path: str, filename: str) -> Dict:
        """
        Process PDF file and add to vector store.
        
        Args:
            file_path: Path to PDF file
            filename: Display name for the file
            
        Returns:
            Dict with processing results
        """
        if not pypdf:
            return {"error": "pypdf not installed. Install with: pip install pypdf"}
        
        if not self.vectorstore:
            return {"error": "Vector store not initialized"}
        
        try:
            # Extract text from PDF
            reader = pypdf.PdfReader(file_path)
            text = ""
            page_count = len(reader.pages)
            
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            if not text.strip():
                return {"error": "No text found in PDF"}
            
            # Split into chunks
            chunks = self._split_text(text)
            
            # Create documents with metadata
            documents = [
                Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "chunk": i,
                        "total_chunks": len(chunks),
                        "page_count": page_count
                    }
                )
                for i, chunk in enumerate(chunks)
            ]
            
            # Add to vector store
            self.vectorstore.add_documents(documents)
            self.vectorstore.save_local(self.persist_dir)
            
            # Update metadata
            self.documents_metadata.append({
                "filename": filename,
                "path": file_path,
                "chunks": len(chunks),
                "pages": page_count
            })
            self._save_metadata()
            
            return {
                "success": True,
                "filename": filename,
                "chunks": len(chunks),
                "pages": page_count
            }
            
        except Exception as e:
            return {"error": f"Failed to process PDF: {str(e)}"}
    
    def process_text(self, file_path: str, filename: str) -> Dict:
        """
        Process text file and add to vector store.
        
        Args:
            file_path: Path to text file
            filename: Display name for the file
            
        Returns:
            Dict with processing results
        """
        if not self.vectorstore:
            return {"error": "Vector store not initialized"}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            if not text.strip():
                return {"error": "Empty file"}
            
            # Split into chunks
            chunks = self._split_text(text)
            
            # Create documents
            documents = [
                Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "chunk": i,
                        "total_chunks": len(chunks)
                    }
                )
                for i, chunk in enumerate(chunks)
            ]
            
            # Add to vector store
            self.vectorstore.add_documents(documents)
            self.vectorstore.save_local(self.persist_dir)
            
            # Update metadata
            self.documents_metadata.append({
                "filename": filename,
                "path": file_path,
                "chunks": len(chunks)
            })
            self._save_metadata()
            
            return {
                "success": True,
                "filename": filename,
                "chunks": len(chunks)
            }
            
        except Exception as e:
            return {"error": f"Failed to process text file: {str(e)}"}
    
    def process_docx(self, file_path: str, filename: str) -> Dict:
        """
        Process DOCX file and add to vector store.
        
        Args:
            file_path: Path to DOCX file
            filename: Display name for the file
            
        Returns:
            Dict with processing results
        """
        if not docx:
            return {"error": "python-docx not installed. Install with: pip install python-docx"}
        
        if not self.vectorstore:
            return {"error": "Vector store not initialized"}
        
        try:
            # Extract text from DOCX
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            if not text.strip():
                return {"error": "No text found in DOCX"}
            
            # Split into chunks
            chunks = self._split_text(text)
            
            # Create documents
            documents = [
                Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "chunk": i,
                        "total_chunks": len(chunks)
                    }
                )
                for i, chunk in enumerate(chunks)
            ]
            
            # Add to vector store
            self.vectorstore.add_documents(documents)
            self.vectorstore.save_local(self.persist_dir)
            
            # Update metadata
            self.documents_metadata.append({
                "filename": filename,
                "path": file_path,
                "chunks": len(chunks)
            })
            self._save_metadata()
            
            return {
                "success": True,
                "filename": filename,
                "chunks": len(chunks)
            }
            
        except Exception as e:
            return {"error": f"Failed to process DOCX: {str(e)}"}
    
    def _split_text(self, text: str) -> List[str]:
        """
        Split text into chunks for embedding.
        
        Args:
            text: Text to split
            
        Returns:
            List of text chunks
        """
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,  # ~250 words
            chunk_overlap=200,  # 50 word overlap
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        return splitter.split_text(text)
    
    def search(self, query: str, k: int = 3) -> List[Dict]:
        """
        Semantic search through indexed documents.
        
        Args:
            query: Search query
            k: Number of results to return
            
        Returns:
            List of results with content, source, and score
        """
        if not self.vectorstore:
            return []
        
        try:
            # Search with scores
            results = self.vectorstore.similarity_search_with_score(query, k=k)
            
            return [
                {
                    "content": doc.page_content,
                    "source": doc.metadata.get("source", "Unknown"),
                    "chunk": doc.metadata.get("chunk", 0),
                    "score": float(score),
                    "relevance": "High" if score < 0.5 else "Medium" if score < 1.0 else "Low"
                }
                for doc, score in results
                if doc.metadata.get("source") != "system"  # Exclude dummy doc
            ]
            
        except Exception as e:
            print(f"❌ Search error: {e}")
            return []
    
    def list_documents(self) -> List[Dict]:
        """
        Get list of indexed documents.
        
        Returns:
            List of document metadata
        """
        return self.documents_metadata
    
    def _save_metadata(self):
        """Save documents metadata to disk"""
        metadata_path = os.path.join(self.persist_dir, "metadata.json")
        with open(metadata_path, 'w') as f:
            json.dump(self.documents_metadata, f, indent=2)
    
    def clear_all(self) -> Dict:
        """
        Clear all documents from vector store.
        
        Returns:
            Dict with operation result
        """
        try:
            # Recreate empty store
            self._create_empty_store()
            self.documents_metadata = []
            self._save_metadata()
            
            return {"success": True, "message": "All documents cleared"}
        except Exception as e:
            return {"error": f"Failed to clear documents: {str(e)}"}


# Global instance (lazy initialization)
_rag_instance = None

def get_rag_system() -> DocumentRAG:
    """Get global RAG system instance"""
    global _rag_instance
    if _rag_instance is None:
        _rag_instance = DocumentRAG()
    return _rag_instance
